from turtle import width
from docx import Document
from docx.shared import Inches
import pyttsx3




def speak(text):
    pyttsx3.speak(text)


document = Document()

# PROFILE PICTURE
document.add_picture(
    'picture.jpg',
     width=Inches(1.89)
)

# PEROSNAL INFORMATION

name = input('what is your name? ')
surname = input('what is your surname? ')
age = input('what is your age? ')
email = input('what is your email? ')
phone_number = input('what is your phone number? ')
nationality = input('what is your nationality? ')
tribe = input('what is your tribe? ')
religion = input('what is your religion? ')
marital_status = input('what is your marital status? ')
profession = input('what is your profession? ')

document.add_paragraph(
    name + ' | ' + surname + ' | ' + str(age) + ' | ' + email + ' | ' + phone_number + ' | ' + nationality + ' | ' + tribe + ' | ' + religion + ' | ' + marital_status + ' | ' + profession  
    )
# ABOUT ME

document.add_heading('OTHER INFO:')
OTHER_INFO = input('Tell me more about yourself please.....  ')
document.add_paragraph(OTHER_INFO)

# WORK EXPERIENCE

document.add_heading('WORK EXPERIENCE')
p = document.add_paragraph()

company = input('Enter your company ')
working_location = input('Enter working location ')
start_date = input('Enter start date  ')
end_date = input('Enter end date  ')
working_type = input('Enter working type here ')

p.add_run(company + ' ' ).bold = True
p.add_run(start_date + '-' + end_date + '\n').italic = True

experience_detils = input('Enter the nature of your experince at ' + company + ' ' ) 
p.add_run(experience_detils)

#MORE WORK EXPERIENCE

while True:
    more_work_experiences = input('Do you have more work experiences?  ')
    if more_work_experiences.upper() == 'YES':
        p = document.add_paragraph()

        company = input('Enter your company ')
        working_location = input('Enter working location ')
        start_date = input('Enter start date  ')
        end_date = input('Enter end date  ')
        working_type = input('Enter working type here ')

        p.add_run(company + ' ' ).bold = True
        p.add_run(start_date + '-' + end_date + '\n').italic = True

        experience_detils = input('Enter the nature of your experince at ' + company + ' ' ) 
        p.add_run(experience_detils)
    else:
        break

# SOFT SKILLS

document.add_heading('SOFT SKILLS')
skills = input('Enter your skills here  ')
p = document.add_paragraph(skills)
p.style = 'List Bullet'

while True:
    relevant_skills = input('Do you have relevant skillls to provide?  ')
    if relevant_skills.lower() == 'yes':
        skills = input('Enter your skills here')
        p = document.add_paragraph(skills) 
        p.style = 'List Bullet'   
    else:
        break

#LANGUAGE SPOKEN

document.add_heading('LANGUAGE SKILLS')
language = input('Enter spoken language here ')
p = document.add_paragraph(language)
p.style = 'List Bullet'

while True:
    language_skills = input('Do you have any spoken language as a skill?  ')
    if language_skills == 'yes':
        language = input('Enter a language here  ')
        p = document.add_paragraph(language)
        p.style = 'List Bullet'
    else:
        break

#FOOTER

section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "A RESUME SUCCESFULLY GENERATED USING PYTHON"
document.save('Resume.docx')
